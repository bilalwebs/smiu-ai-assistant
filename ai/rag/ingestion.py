"""Document ingestion and chunking (AI_ARCHITECTURE.md §14.2, §36).

Purpose:
    The deterministic, offline front half of the RAG pipeline: read a source
    document, validate it, split it into retrievable chunks with source
    metadata, and produce the in-memory document/chunk records that later
    phases persist and embed. Nothing in this module calls a model, a vector
    store, or the network (AI_ARCHITECTURE.md §36.2-36.5).

Lifecycle (AI_ARCHITECTURE.md §36.2):
    Read → extract text → validate → chunk → embed → index

    This module implements the first four stages: reading the source text,
    extracting it per file type, validating it (file type, category, version,
    checksum, content present — §36.3), and chunking it deterministically
    (§36.4). Embedding (§15) and indexing are later tasks in Phase 9.

Validation rules (AI_ARCHITECTURE.md §36.3):
    - File type whitelist: pdf / md / txt / docx.
    - SHA-256 checksum over the canonical (CRLF-normalized) text.
    - Extractable text must be present; binary-only files are rejected.
    - Category must map to a knowledge category (admission / examination /
      faq / documents).

Chunking rules (AI_ARCHITECTURE.md §36.4):
    - Semantic/size-bounded chunks; paragraphs are the unit and markdown
      headings attach section context to every following chunk.
    - A small overlap preserves context across split boundaries — only
      within a single oversized paragraph, so chunks are never duplicated.
    - Deterministic: same input → same chunks, same chunk ids.
    - Every chunk carries heading/page/token/character metadata (§36.5,
      DATABASE_DESIGN.md §21.2).
"""

from __future__ import annotations

import hashlib
import importlib.util
import io
import re
import uuid
from collections.abc import Iterable, Sequence
from enum import StrEnum
from pathlib import Path
from typing import Any

from pydantic import BaseModel, Field

from ai.core.state import RetrievedChunk


class IngestionError(Exception):
    """Base class for all document-ingestion failures (§36)."""


class UnsupportedFileTypeError(IngestionError):
    """The file extension is not on the §36.3 whitelist (pdf/md/txt/docx)."""


class EmptyDocumentError(IngestionError):
    """The source contains no extractable text (§36.3 content scan)."""


class MalformedDocumentError(IngestionError):
    """The source text cannot be decoded as the declared file type."""


class InvalidMetadataError(IngestionError):
    """Document header/folder metadata is missing, unknown, or invalid."""


class DuplicateSourceError(IngestionError):
    """The same (source_path, version) already exists in this ingest batch."""


class DocumentParseError(IngestionError):
    """The optional parser for a binary file type is unavailable or failed.

    pdf/docx extraction needs the ``pypdf`` / ``python-docx`` packages. The
    ingestion pipeline never silently drops or mis-reads a binary source:
    it fails loudly with this typed error so the caller can choose to install
    the parser or reject the upload (AI_ARCHITECTURE.md §36.3).
    """


class KnowledgeCategory(StrEnum):
    """Knowledge categories (DATABASE_DESIGN.md §21.1)."""

    ADMISSION = "admission"
    EXAMINATION = "examination"
    FAQ = "faq"
    DOCUMENTS = "documents"


KNOWN_CATEGORY_FOLDERS = tuple(item.value for item in KnowledgeCategory)

SUPPORTED_FILE_TYPES: dict[str, str] = {
    ".md": "md",
    ".txt": "txt",
    ".pdf": "pdf",
    ".docx": "docx",
}

_MARKDOWN_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")
_VERSION_RE = re.compile(r"^[0-9]+(?:\.[0-9]+)*$")
_MAX_VERSION_LENGTH = 30
_MAX_CHUNK_ID_LENGTH = 64


class DocumentChunk(BaseModel):
    """One retrievable chunk plus source metadata (§36.4-36.5, §21.2)."""

    chunk_id: str = Field(min_length=1, max_length=_MAX_CHUNK_ID_LENGTH)
    document_id: uuid.UUID | None = None
    title: str = Field(min_length=1, max_length=255)
    category: str = Field(min_length=1)
    version: str = Field(min_length=1, max_length=_MAX_VERSION_LENGTH)
    source_path: str = ""
    chunk_index: int = Field(default=0, ge=0)
    heading: str | None = None
    page_number: int | None = None
    chunk_text: str = Field(min_length=1)
    token_count: int = Field(ge=1)
    character_count: int = Field(ge=1)
    checksum_sha256: str = Field(min_length=64, max_length=64)

    def to_retrieved_chunk(self, *, score: float = 0.0) -> RetrievedChunk:
        """Map this chunk to the retrieval contract (ai/rag/retriever.py).

        The snippet is the full chunk text so the context builder and the
        citation assembler consume the same unit the index will store.
        """
        return RetrievedChunk(
            chunk_id=self.chunk_id,
            document_id=self.document_id,
            title=self.title,
            category=self.category,
            snippet=self.chunk_text,
            score=score,
        )


class IngestedDocument(BaseModel):
    """An ingested knowledge document plus its chunks (§21.1, §36.5).

    ``status`` defaults to ``processed`` and ``is_active`` to ``True`` so an
    ingested document is immediately eligible for retrieval; the backend
    persistence layer maps this record to ``knowledge_documents`` and the
    retriever's candidate filter (DATABASE_DESIGN.md §21.3).
    """

    document_id: uuid.UUID = Field(default_factory=uuid.uuid4)
    title: str = Field(min_length=1, max_length=255)
    category: str = Field(min_length=1)
    source_path: str = ""
    file_type: str | None = None
    file_size: int | None = Field(default=None, ge=0)
    author: str | None = Field(default=None, max_length=150)
    version: str = Field(default="1", min_length=1, max_length=_MAX_VERSION_LENGTH)
    checksum_sha256: str = Field(min_length=64, max_length=64)
    status: str = "processed"
    chunk_count: int = Field(ge=0)
    is_active: bool = True
    metadata: dict[str, Any] = Field(default_factory=dict)
    chunks: list[DocumentChunk] = Field(default_factory=list)

    def is_eligible(self, *, current_version: str | None = None) -> bool:
        """Whether this document is a retrieval candidate (§21.3, §16.4).

        Only ``is_active`` + ``status='processed'`` documents participate,
        and only when no newer version is current (current version only).
        """
        return (
            self.is_active
            and self.status == "processed"
            and (current_version is None or self.version == current_version)
        )


def checksum_sha256(text: str) -> str:
    """Return the SHA-256 hex digest of the canonical document text."""
    return hashlib.sha256(normalize_text(text).encode("utf-8")).hexdigest()


def normalize_text(text: str) -> str:
    """Canonicalize line endings so checksums are platform-stable (§36.3)."""
    return text.replace("\r\n", "\n").replace("\r", "\n")


def validate_title(title: str) -> str:
    """Validate and return a non-empty, stripped document title."""
    value = title.strip()
    if not value:
        raise InvalidMetadataError("title must not be empty")
    if len(value) > 255:
        raise InvalidMetadataError("title must be at most 255 characters")
    return value


def validate_category(category: str) -> str:
    """Coerce a category to a known knowledge category (§36.3)."""
    try:
        return KnowledgeCategory(category).value
    except ValueError as exc:
        raise InvalidMetadataError(
            f"unknown category {category!r}; expected one of "
            f"{', '.join(KNOWN_CATEGORY_FOLDERS)}"
        ) from exc


def validate_version(version: str) -> str:
    """Validate a version string (§36.5; e.g. ``1``, ``1.2``, ``2.0.1``)."""
    value = version.strip()
    if not value:
        raise InvalidMetadataError("version must not be empty")
    if len(value) > _MAX_VERSION_LENGTH:
        raise InvalidMetadataError(
            f"version must be at most {_MAX_VERSION_LENGTH} characters"
        )
    if not _VERSION_RE.match(value):
        raise InvalidMetadataError(
            f"invalid version {value!r}; expected dotted integers like '1' or '1.2'"
        )
    return value


def _stable_chunk_id(
    *, source_path: str, version: str, chunk_index: int
) -> str:
    """Return a deterministic chunk id (sha256 prefix) for a chunk position.

    Same source + version + position always yields the same id, so repeated
    ingestion is idempotent at the id level and re-indexing is reproducible
    (AI_ARCHITECTURE.md §36.7 idempotency).
    """
    raw = f"{source_path}\x1f{version}\x1f{chunk_index}"
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()[:_MAX_CHUNK_ID_LENGTH]


def _estimate_tokens(text: str) -> int:
    """Approximate English tokens-per-character ratio (matches §17).

    This mirrors ``ai.rag.context_builder._default_estimator`` so ingestion
    and context budgeting agree on chunk sizes.
    """
    return max(1, len(text) // 4)


def _split_paragraphs(text: str) -> list[str]:
    """Split normalized text into non-empty paragraphs (blank-line separated)."""
    return [part.strip() for part in text.split("\n\n") if part.strip()]


def _split_sentences(text: str) -> list[str]:
    """Split a paragraph into sentences on sentence-final punctuation."""
    parts = re.split(r"(?<=[.!?])\s+", text.strip())
    return [part for part in parts if part]


def chunk_text(
    *,
    text: str,
    chunk_size: int = 800,
    overlap: int = 80,
) -> list[tuple[str, str | None]]:
    """Split ``text`` into size-bounded chunks (paragraph unit, §36.4).

    Returns ``(chunk_text, heading)`` pairs in document order. Chunking rules:
    - paragraphs are never split across chunks unless a single paragraph
      exceeds ``chunk_size``;
    - oversized paragraphs are split on sentence boundaries with ``overlap``
      tail context (overlap only within the same paragraph, so no text is
      duplicated across chunk boundaries);
    - the most recent markdown heading attaches to every following chunk;
    - the result is deterministic: same input → same chunk sequence.
    """
    if chunk_size < 1:
        raise ValueError("chunk_size must be >= 1")
    if overlap < 0:
        raise ValueError("overlap must be >= 0")

    current_heading: str | None = None
    chunks: list[tuple[str, str | None]] = []

    for paragraph in _split_paragraphs(normalize_text(text)):
        match = _MARKDOWN_HEADING_RE.match(paragraph)
        if match:
            current_heading = match.group(2).strip()
            continue

        if len(paragraph) <= chunk_size:
            chunks.append((paragraph, current_heading))
            continue

        for piece in _split_paragraph(paragraph, chunk_size, overlap):
            chunks.append((piece, current_heading))

    return chunks


def _split_paragraph(paragraph: str, chunk_size: int, overlap: int) -> list[str]:
    """Split one oversized paragraph on sentence boundaries with overlap."""
    pieces: list[str] = []
    buffer = ""
    for sentence in _split_sentences(paragraph):
        candidate = f"{buffer} {sentence}".strip() if buffer else sentence
        if len(candidate) <= chunk_size:
            buffer = candidate
            continue
        if buffer:
            pieces.append(buffer)
        if len(sentence) > chunk_size:
            for hard in _hard_split(sentence, chunk_size):
                pieces.append(hard)
            buffer = ""
            continue
        buffer = sentence

    if buffer:
        pieces.append(buffer)

    if len(pieces) <= 1:
        return pieces

    overlapped: list[str] = []
    for index, piece in enumerate(pieces):
        overlapped.append(_with_overlap(piece, pieces, index, overlap))
    return overlapped


def _hard_split(text: str, chunk_size: int) -> list[str]:
    """Split a single sentence longer than ``chunk_size`` on word bounds."""
    words = text.split()
    pieces: list[str] = []
    buffer: list[str] = []
    buffer_len = 0
    for word in words:
        if buffer and buffer_len + len(word) + 1 > chunk_size:
            pieces.append(" ".join(buffer))
            buffer = [word]
            buffer_len = len(word)
        else:
            buffer.append(word)
            buffer_len += len(word) + 1 if len(buffer) > 1 else len(word)
    if buffer:
        pieces.append(" ".join(buffer))
    return pieces or [text]


def _with_overlap(
    piece: str, all_pieces: Sequence[str], index: int, overlap: int
) -> str:
    """Attach the previous piece's tail as overlap context (if any)."""
    if index == 0 or overlap <= 0:
        return piece
    previous = all_pieces[index - 1]
    tail = _overlap_tail(previous, overlap)
    if not tail:
        return piece
    return f"{tail} {piece}"


def _overlap_tail(text: str, overlap: int) -> str:
    """Return up to ``overlap`` characters of the end of ``text``.

    The tail is cut on a word boundary and kept non-empty only when it is
    long enough to be useful.
    """
    if len(text) <= overlap:
        return ""
    cut = text[-overlap:]
    boundary = cut.find(" ")
    if boundary > 0:
        cut = cut[boundary + 1 :]
    if len(cut) < 8:
        return ""
    return cut.strip()


def _file_type_from_path(path: str) -> str | None:
    return SUPPORTED_FILE_TYPES.get(Path(path).suffix.lower())


def extract_text(path: str, *, data: bytes | None = None) -> str:
    """Extract document text from ``path`` (or ``data`` when provided).

    - ``txt``/``md``: UTF-8, strict decode (a decode failure raises
      ``MalformedDocumentError`` — no silent mojibake).
    - ``pdf``/``docx``: extracted through the optional ``pypdf`` /
      ``python-docx`` parsers. When a parser package is not installed the
      call fails loudly with ``DocumentParseError`` (never a silent skip);
      when it is installed the parsed text is used.
    - Returns the canonical (normalized) text. Blank/whitespace-only sources
      raise ``EmptyDocumentError`` (§36.3 content scan).
    """
    file_type = _file_type_from_path(path)
    if file_type is None:
        raise UnsupportedFileTypeError(
            f"unsupported file type for {path!r}; supported: "
            f"{', '.join(sorted(SUPPORTED_FILE_TYPES))}"
        )

    if file_type == "txt" or file_type == "md":
        raw = Path(path).read_bytes() if data is None else data
        try:
            text = raw.decode("utf-8")
        except UnicodeDecodeError as exc:
            raise MalformedDocumentError(
                f"document {path!r} is not valid UTF-8 text"
            ) from exc
    elif file_type == "pdf":
        text = _extract_pdf(path, data)
    elif file_type == "docx":
        text = _extract_docx(path, data)
    else:  # pragma: no cover - _file_type_from_path guards this
        raise UnsupportedFileTypeError(f"unsupported file type for {path!r}")

    normalized = normalize_text(text)
    if not normalized.strip():
        raise EmptyDocumentError(f"document {path!r} contains no extractable text")
    return normalized


def _extract_pdf(path: str, data: bytes | None) -> str:
    """Extract text from a PDF via the optional ``pypdf`` parser."""
    if importlib.util.find_spec("pypdf") is None:
        raise DocumentParseError(
            "PDF extraction requires the optional 'pypdf' package; "
            f"cannot parse {path!r}"
        )
    try:
        from pypdf import PdfReader  # type: ignore[import-not-found]
    except ImportError as exc:  # pragma: no cover - find_spec guards this
        raise DocumentParseError(
            f"PDF parser unavailable for {path!r}"
        ) from exc

    reader = PdfReader(path if data is None else io.BytesIO(data))
    pages = [" ".join((page.extract_text() or "").split()) for page in reader.pages]
    return "\n".join(page for page in pages if page)


def _extract_docx(path: str, data: bytes | None) -> str:
    """Extract text from a .docx via the optional ``python-docx`` parser."""
    if importlib.util.find_spec("docx") is None:
        raise DocumentParseError(
            "docx extraction requires the optional 'python-docx' package; "
            f"cannot parse {path!r}"
        )
    try:
        import docx  # type: ignore[import-not-found]
    except ImportError as exc:  # pragma: no cover - find_spec guards this
        raise DocumentParseError(
            f"docx parser unavailable for {path!r}"
        ) from exc

    document = docx.Document(path if data is None else io.BytesIO(data))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


_FRONT_MATTER_DELIMITER = "---"
_FRONT_MATTER_KEYS = ("title", "category", "version", "author")


def parse_document_header(text: str) -> tuple[dict[str, str], str]:
    """Parse an optional markdown front-matter header (§36.5).

    Returns ``(metadata, body)``. A leading ``---`` block may declare
    ``title`` / ``category`` / ``version`` / ``author``. Unknown keys,
    duplicate keys, and empty values are rejected with
    ``InvalidMetadataError``. A document without a header returns an empty
    metadata dict and the full text as the body.
    """
    normalized = normalize_text(text)
    lines = normalized.splitlines()
    if len(lines) < 3 or lines[0].strip() != _FRONT_MATTER_DELIMITER:
        return {}, text

    metadata: dict[str, str] = {}
    for line_index, line in enumerate(lines[1:], start=1):
        stripped = line.strip()
        if stripped == _FRONT_MATTER_DELIMITER:
            return metadata, "\n".join(lines[line_index + 1 :]).strip()
        if not stripped or stripped.startswith("#"):
            raise InvalidMetadataError(
                "front matter entries must be 'key: value' pairs"
            )
        if ":" not in stripped:
            raise InvalidMetadataError(
                f"front matter entry {stripped!r} is missing a ': value'"
            )
        key, _, value = stripped.partition(":")
        key = key.strip().lower()
        value = value.strip()
        if key not in _FRONT_MATTER_KEYS:
            raise InvalidMetadataError(
                f"unknown front matter key {key!r}; allowed: "
                f"{', '.join(_FRONT_MATTER_KEYS)}"
            )
        if key in metadata:
            raise InvalidMetadataError(f"duplicate front matter key {key!r}")
        if not value:
            raise InvalidMetadataError(f"front matter key {key!r} has an empty value")
        metadata[key] = value

    raise InvalidMetadataError(
        "unterminated front matter block (missing closing '---')"
    )


def ingest_document(
    *,
    text: str,
    title: str,
    category: str,
    version: str = "1",
    source_path: str = "",
    author: str | None = None,
    file_type: str | None = None,
    file_size: int | None = None,
    chunk_size: int = 800,
    overlap: int = 80,
    document_id: uuid.UUID | None = None,
    checksum: str | None = None,
) -> IngestedDocument:
    """Validate, chunk, and package one document (ingest stage, §14.2/§36)."""
    if not text.strip():
        raise EmptyDocumentError("document text must not be empty")
    if file_type is not None and file_type not in SUPPORTED_FILE_TYPES.values():
        raise UnsupportedFileTypeError(
            f"unsupported file type {file_type!r}; supported: "
            f"{', '.join(sorted(SUPPORTED_FILE_TYPES.values()))}"
        )
    if file_size is not None and file_size < 0:
        raise ValueError("file_size must be >= 0")

    normalized = normalize_text(text)
    canonical_title = validate_title(title)
    canonical_category = validate_category(category)
    canonical_version = validate_version(version)
    document_checksum = checksum or checksum_sha256(normalized)
    if len(document_checksum) != 64:
        raise ValueError("checksum must be a 64-character SHA-256 hex digest")

    document_id_value = document_id or uuid.uuid4()
    chunks: list[DocumentChunk] = []
    for index, (chunk_text_value, heading) in enumerate(
        chunk_text(text=normalized, chunk_size=chunk_size, overlap=overlap)
    ):
        chunks.append(
            DocumentChunk(
                chunk_id=_stable_chunk_id(
                    source_path=source_path,
                    version=canonical_version,
                    chunk_index=index,
                ),
                document_id=document_id_value,
                title=canonical_title,
                category=canonical_category,
                version=canonical_version,
                source_path=source_path,
                chunk_index=index,
                heading=heading,
                chunk_text=chunk_text_value,
                token_count=_estimate_tokens(chunk_text_value),
                character_count=len(chunk_text_value),
                checksum_sha256=document_checksum,
            )
        )

    return IngestedDocument(
        document_id=document_id_value,
        title=canonical_title,
        category=canonical_category,
        source_path=source_path,
        file_type=file_type,
        file_size=file_size,
        author=author,
        version=canonical_version,
        checksum_sha256=document_checksum,
        status="processed",
        chunk_count=len(chunks),
        is_active=True,
        metadata={
            "author": author,
            "file_type": file_type,
            "file_size": file_size,
        },
        chunks=chunks,
    )


def ingest_documents(
    documents: Iterable[dict[str, Any]],
    *,
    chunk_size: int = 800,
    overlap: int = 80,
) -> list[IngestedDocument]:
    """Ingest many documents, rejecting duplicate ``(source_path, version)``.

    The partial-unique ``(source_path, version)`` constraint
    (DATABASE_DESIGN.md §21.1) is enforced in memory at the batch boundary;
    two documents with the same source path and version raise
    ``DuplicateSourceError``.
    """
    seen: set[tuple[str, str]] = set()
    ingested: list[IngestedDocument] = []
    for raw in documents:
        source_path = str(raw.get("source_path") or "")
        version = validate_version(str(raw.get("version") or "1"))
        key = (source_path, version)
        if key in seen:
            raise DuplicateSourceError(
                f"duplicate document source (source_path={source_path!r}, "
                f"version={version!r}) in ingest batch"
            )
        seen.add(key)
        ingested.append(
            ingest_document(
                text=str(raw["text"]),
                title=str(raw["title"]),
                category=str(raw["category"]),
                version=version,
                source_path=source_path,
                author=raw.get("author"),
                file_type=raw.get("file_type"),
                file_size=raw.get("file_size"),
                chunk_size=chunk_size,
                overlap=overlap,
                document_id=raw.get("document_id"),
                checksum=raw.get("checksum"),
            )
        )
    return ingested


class KnowledgeIngestor:
    """Directory-oriented ingestion over the ``knowledge/`` root (§36.2).

    Discovers only supported files inside category folders
    (admission / examination / faq / documents), skips hidden dotfiles such
    as ``.gitkeep``, and yields validated, chunked ``IngestedDocument``
    records. Parsing of pdf/docx still requires their optional packages
    (see ``DocumentParseError``).
    """

    def __init__(
        self,
        knowledge_root: str | Path,
        *,
        chunk_size: int = 800,
        overlap: int = 80,
    ) -> None:
        self.knowledge_root = Path(knowledge_root)
        self.chunk_size = chunk_size
        self.overlap = overlap

    def discover(self) -> list[Path]:
        """Return supported document files under category folders, in order."""
        found: list[Path] = []
        for folder in KNOWN_CATEGORY_FOLDERS:
            folder_path = self.knowledge_root / folder
            if not folder_path.is_dir():
                continue
            for candidate in sorted(folder_path.rglob("*")):
                if not candidate.is_file():
                    continue
                if any(part.startswith(".") for part in candidate.parts):
                    continue
                if _file_type_from_path(candidate.name) is None:
                    continue
                found.append(candidate)
        return found

    def ingest_directory(self) -> list[IngestedDocument]:
        """Ingest every discovered document in the knowledge root (§36.2)."""
        documents: list[dict[str, Any]] = []
        for path in self.discover():
            text = extract_text(str(path))
            header, body = parse_document_header(text)
            relative = path.relative_to(self.knowledge_root).as_posix()
            category = _category_from_folder(path)
            documents.append(
                {
                    "text": body or text,
                    "title": header.get("title") or path.stem,
                    "category": header.get("category", category),
                    "version": header.get("version", "1"),
                    "author": header.get("author"),
                    "source_path": relative,
                    "file_type": _file_type_from_path(path.name),
                    "file_size": path.stat().st_size,
                }
            )
        return ingest_documents(documents, chunk_size=self.chunk_size, overlap=self.overlap)


def _category_from_folder(path: Path) -> str:
    """Derive the knowledge category from the document's folder (§36.5)."""
    relative = path.as_posix().split("/")
    for folder in KNOWN_CATEGORY_FOLDERS:
        if folder in relative:
            return folder
    raise InvalidMetadataError(
        f"document {path!r} is not inside a known knowledge category folder "
        f"({', '.join(KNOWN_CATEGORY_FOLDERS)})"
    )
